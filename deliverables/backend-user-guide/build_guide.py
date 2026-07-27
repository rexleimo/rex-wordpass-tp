"""Build the customer-facing WordPress backend operating guide."""

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\coding\rex-wordpass-tp\deliverables\backend-user-guide")
SOURCE = ROOT / "screenshots"
ANNOTATED = ROOT / "annotated"
OUTPUT = ROOT / "toKraft-WordPress后台使用指南-截图复核版.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
GOLD = "B17B20"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
MUTED = "5B6573"
RED = "9B1C1C"
FONT = "Microsoft YaHei"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_paragraph(paragraph, before=0, after=6, line=1.25, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def add_text(doc, text, size=10.5, color="000000", bold=False, before=0, after=6, align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph(p, before, after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph(hp, after=0, line=1.0)
    run = hp.add_run("toKraft | WordPress 后台使用指南")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(fp, before=0, after=0, line=1.0)
    run = fp.add_run("客户交付版本  |  第 ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(fp)
    run = fp.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def safe_filename(name):
    return ANNOTATED / name


def load_font(size):
    for path in [r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\arial.ttf"]:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def annotate(source_name, output_name, markers=(), content_box=None, boxes=()):
    image = Image.open(SOURCE / source_name).convert("RGB")
    if content_box:
        image = image.crop(content_box)
    overlay = Image.new("RGBA", image.size, (0, 0, 0, 0))
    overlay_draw = ImageDraw.Draw(overlay)
    for number, x1, y1, x2, y2 in boxes:
        # Region boxes are tied to the freshly captured viewport, not guessed offsets.
        overlay_draw.rectangle((x1, y1, x2, y2), fill=(177, 123, 32, 35), outline=(177, 123, 32, 255), width=4)
        overlay_draw.ellipse((x1 - 22, y1 - 22, x1 + 22, y1 + 22), fill=(177, 123, 32, 255), outline="white", width=3)
        label_font = load_font(22)
        label_box = overlay_draw.textbbox((0, 0), str(number), font=label_font)
        overlay_draw.text((x1 - (label_box[2] - label_box[0]) / 2, y1 - (label_box[3] - label_box[1]) / 2 - 2), str(number), fill="white", font=label_font)
    image = Image.alpha_composite(image.convert("RGBA"), overlay).convert("RGB")
    draw = ImageDraw.Draw(image)
    font = load_font(24)
    for number, x, y in markers:
        radius = 23
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill="#B17B20", outline="white", width=3)
        bbox = draw.textbbox((0, 0), str(number), font=font)
        tx = x - (bbox[2] - bbox[0]) / 2
        ty = y - (bbox[3] - bbox[1]) / 2 - 2
        draw.text((tx, ty), str(number), fill="white", font=font)
    out = safe_filename(output_name)
    out.parent.mkdir(parents=True, exist_ok=True)
    image.save(out, quality=92)
    return out


def add_legend(doc, items):
    p = doc.add_paragraph()
    set_paragraph(p, before=4, after=4, line=1.15)
    for index, item in enumerate(items):
        run = p.add_run(f"{index + 1}. {item}")
        set_run_font(run, size=9.2, color=NAVY, bold=True)
        if index < len(items) - 1:
            separator = p.add_run("     ")
            set_run_font(separator, size=9.2)


def add_note_box(doc, label, text, color=BLUE, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph(p, before=0, after=0, line=1.2)
    run = p.add_run(f"{label}  ")
    set_run_font(run, size=10, color=color, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, before=0, after=4, line=1.25)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="000000")


def add_figure_page(doc, heading, intro, image_path, legend, steps, safety=None):
    doc.add_page_break()
    add_heading(doc, heading, 1)
    add_text(doc, intro, size=10.4, color=MUTED, before=0, after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=0, after=0, line=1.0)
    p.add_run().add_picture(str(image_path), width=Inches(6.15))
    add_legend(doc, legend)
    for item in steps:
        add_bullet(doc, item)
    if safety:
        add_note_box(doc, "操作提示", safety, GOLD, "FFF8E8")


def add_capture_page(doc, heading, intro, image_path, steps, safety=None):
    """Add a full-page capture without artificial markers when the UI is self-explanatory."""
    doc.add_page_break()
    add_heading(doc, heading, 1)
    add_text(doc, intro, size=10.4, color=MUTED, before=0, after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=0, after=4, line=1.0)
    with Image.open(image_path) as capture:
        height_at_standard_width = Inches(6.15).emu * capture.height / capture.width
    if height_at_standard_width > Inches(6.35).emu:
        p.add_run().add_picture(str(image_path), height=Inches(6.35))
    else:
        p.add_run().add_picture(str(image_path), width=Inches(6.15))
    for item in steps:
        add_bullet(doc, item)
    if safety:
        add_note_box(doc, "操作提示", safety, GOLD, "FFF8E8")


def build():
    ANNOTATED.mkdir(parents=True, exist_ok=True)
    dashboard = annotate("01-dashboard.png", "01-dashboard-annotated.png", boxes=[(1, 0, 38, 157, 710), (2, 180, 550, 678, 735), (3, 692, 550, 1202, 946)])
    posts = annotate("10-blog-posts.png", "10-blog-posts-annotated.png", boxes=[(1, 232, 51, 302, 83), (2, 180, 96, 462, 122), (3, 180, 180, 1230, 486), (4, 0, 88, 157, 240)])
    products = annotate("03-products-list.png", "03-products-list-annotated.png", boxes=[(1, 232, 110, 322, 144), (2, 180, 196, 930, 232), (3, 180, 269, 1230, 1238), (4, 0, 302, 157, 543)])
    orders = annotate("09-shop-orders.png", "09-shop-orders-annotated.png", boxes=[(1, 234, 110, 315, 144), (2, 180, 199, 733, 233), (3, 180, 245, 1230, 451), (4, 0, 276, 157, 528)])
    home_operations = annotate("02-home-operations.png", "02-home-operations-annotated.png", boxes=[(1, 180, 60, 1200, 509), (2, 180, 535, 1200, 1235)])
    product_edit = annotate("04a-product-price.png", "04-product-price-annotated.png", boxes=[(1, 180, 565, 904, 628), (2, 482, 640, 705, 683), (3, 1128, 742, 1194, 786), (4, 925, 828, 1200, 1215)])
    product_stock = annotate("04b-product-stock.png", "04-product-stock-annotated.png", boxes=[(1, 180, 628, 324, 667), (2, 482, 599, 705, 643), (3, 482, 718, 705, 741), (4, 482, 764, 705, 856)])
    quote_records = annotate("05-quote-records.png", "05-quote-records-annotated.png", boxes=[(1, 180, 175, 1230, 586), (2, 223, 219, 350, 500), (3, 655, 219, 736, 500), (4, 0, 234, 157, 401)])
    quote_detail = annotate("06-quote-detail.png", "06-quote-detail-annotated.png", boxes=[(1, 182, 158, 931, 672), (2, 415, 682, 497, 726), (3, 415, 752, 497, 796), (4, 414, 821, 905, 1013), (5, 1160, 259, 1222, 303)])
    materials = annotate("07-material-library.png", "07-material-library-annotated.png", boxes=[(1, 180, 145, 520, 1260), (2, 543, 180, 1200, 598)])
    payments = annotate("08-payments-settings.png", "08-payments-settings-annotated.png", boxes=[(1, 180, 98, 671, 134), (2, 208, 224, 1200, 360), (3, 208, 383, 1200, 469), (4, 208, 488, 1200, 566)])
    blog_edit = annotate("11-blog-edit.png", "11-blog-edit-annotated.png", boxes=[(1, 55, 129, 901, 211), (2, 55, 240, 901, 1240), (3, 970, 111, 1235, 757), (4, 1152, 15, 1204, 49)])

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [(1, 16, BLUE, 18, 10), (2, 13, BLUE, 14, 7), (3, 12, NAVY, 10, 5)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Customer-pack cover.
    add_text(doc, "客户交付资料", size=10, color=GOLD, bold=True, before=30, after=8)
    title = doc.add_paragraph()
    set_paragraph(title, before=0, after=8, line=1.0)
    title_run = title.add_run("toKraft 网站后台使用指南")
    set_run_font(title_run, size=28, color=NAVY, bold=True)
    add_text(doc, "WordPress + WooCommerce 内容、产品与订单日常操作手册", size=13, color=MUTED, after=24)

    meta = doc.add_table(rows=2, cols=2)
    set_table_geometry(meta, [4680, 4680])
    metadata = [("适用范围", "内容维护、产品维护、订单处理、基础站点管理"), ("交付版本", "V1.0 | 2026 年 7 月 25 日"), ("后台入口", "http://localhost:8080/wp-admin/"), ("建议角色", "站点运营、内容编辑、商城运营"),]
    for cell, (label, value) in zip([cell for row in meta.rows for cell in row.cells], metadata):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.15)
        r = p.add_run(label + "\n")
        set_run_font(r, size=9, color=BLUE, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=NAVY)

    add_text(doc, "本指南基于当前已配置的实际后台界面整理。文中带有编号的截图用于定位功能入口；菜单名称可能会随着插件或权限调整而略有变化。", size=10.5, color=NAVY, before=24, after=14)
    add_note_box(doc, "开始前", "请使用已授权账号登录后台。编辑前先确认当前环境；若是正式站点，建议先完成备份或在测试环境验证。", GOLD, "FFF8E8")

    doc.add_page_break()
    add_heading(doc, "1. 使用边界与后台导航", 1)
    add_text(doc, "后台主要分为内容、媒体、页面、WooCommerce、产品、支付和外观设置等模块。日常维护应遵循“先预览、后发布；先小范围验证、再批量调整”的原则。", size=10.5, color="000000")
    add_heading(doc, "建议的权限分工", 2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1800, 3300, 4260])
    headers = ["角色", "可处理事项", "避免直接修改"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    rows = [
        ("内容编辑", "文章、图片、页面文案", "站点地址、插件、支付参数"),
        ("商城运营", "产品、库存、订单状态", "税率、物流、支付密钥"),
        ("管理员", "用户、菜单、外观、系统设置", "未经确认的批量删除或迁移"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=10, color="000000")
    add_heading(doc, "每日操作清单", 2)
    for text in [
        "先检查仪表盘中的系统提醒、更新提示与 WooCommerce 待办。",
        "内容或产品改动完成后，在前台页面用无痕窗口复核链接、图片和价格。",
        "订单处理前核对状态、客户信息和付款情况；不要以测试订单替代真实交易记录。",
    ]:
        add_bullet(doc, text)
    add_note_box(doc, "高风险设置", "不要在“设置 > 常规”中随意修改 WordPress 地址或站点地址；这会造成本地、临时域名或正式域名访问和登录异常。", RED, "FDECEC")

    add_figure_page(
        doc,
        "2. 仪表盘：从哪里开始",
        "仪表盘汇总站点更新、WooCommerce 待办及快捷入口。每次进入后台建议先查看这里，再进入具体模块。",
        dashboard,
        ["左侧主菜单", "页面标题与后台概览", "WooCommerce 设置卡片"],
        [
            "通过 1 号区域切换到文章、媒体、页面、产品和 WooCommerce 模块。",
            "在 2 号区域确认自己所在的后台页面，避免误把内容更新到错误模块。",
            "如 3 号区域出现设置提醒，应由管理员确认后再进行配置，不建议内容编辑人员直接修改。",
        ],
        "如页面出现更新提示，先记录插件或主题名称，再在维护窗口中统一处理。",
    )

    add_figure_page(
        doc,
        "3. 文章：发布 Blog 与案例内容",
        "文章模块用于维护 Blog、指南和案例类内容。当前后台已包含实际文章数据及分类信息。",
        posts,
        ["写文章", "状态与筛选", "文章标题列表", "文章菜单"],
        [
            "点击 1 号“写文章”新建内容；先填写标题、正文和特色图片，再补充分组与标签。",
            "使用 2 号状态标签检查“已发布 / 草稿”，需要复核的内容先保存为草稿。",
            "在 3 号标题列表中点击现有文章即可进入编辑；修改后先预览，再点击更新。",
            "4 号菜单下可维护文章分类与标签，建议沿用已有命名体系，避免重复分类。",
        ],
        "特色图片建议使用统一比例且压缩后再上传；发布后请在前台确认封面、标题换行和正文图片是否正常。",
    )

    add_figure_page(
        doc,
        "4. 产品：维护商品、价格与库存",
        "产品模块承载商城商品。列表中可以快速查看产品图片、SKU、库存、价格和分类。",
        products,
        ["添加新产品", "筛选条件", "产品数据行", "产品菜单"],
        [
            "通过 1 号“添加新产品”创建商品；完成名称、描述、图片、价格、库存和分类后再发布。",
            "使用 2 号区域按分类、产品类型或库存状态筛选，以便做批量核对。",
            "点击 3 号产品名称进入编辑页；SKU 应保持唯一，改价和库存调整需复核。",
            "4 号菜单提供分类、标签、属性及材料库入口；属性命名应保持一致，避免同义重复。",
        ],
        "批量操作前先在筛选结果中确认条目数。若商品涉及变体、税费或库存联动，请只由商城运营或管理员处理。",
    )

    add_figure_page(
        doc,
        "5. 订单：查看、核对与跟进", 
        "订单由 WooCommerce 统一管理。当前订单列表支持按日期、销售渠道和客户条件进行筛选。",
        orders,
        ["添加订单", "订单编号", "订单状态", "WooCommerce 订单菜单"],
        [
            "日常处理应从 4 号 WooCommerce > 订单进入，不建议通过浏览器收藏的旧地址直接操作。",
            "点击 2 号订单编号进入详情，核对商品、金额、客户信息和订单备注。",
            "根据实际处理进度更新 3 号状态；状态变化可能触发邮件或库存动作，提交前需确认。",
            "1 号“添加订单”通常仅用于人工创建订单或内部流程；正常线上订单会自动进入列表。",
        ],
        "不要删除或伪造订单来测试流程。测试请使用独立测试订单，并在完成验证后按既定规则归档。",
    )

    add_capture_page(
        doc,
        "6. 首页运营：内容区块编辑器",
        "入口为 toKraft > 首页内容区块。此页面直接维护首屏、服务卡、设备、材料与底部说明等运营字段。",
        home_operations,
        [
            "先在顶部的内容入口地图确认要改的是首页文案、产品、案例、材料、设备、Blog 或订单。",
            "文本、按钮链接、展示数量和图片均可在对应区块编辑；保存前核对链接使用站内相对路径，例如 /quote/ 或 /shop/。",
            "轮播图建议使用统一的正方形素材；服务卡建议使用统一的 4:5 竖图，避免前台裁切不一致。",
        ],
    )

    add_capture_page(
        doc,
        "7. 商品编辑：图片、SKU、价格与库存入口",
        "单个商品页面包含商品图片、分类、材料关联，以及“产品数据”中的常规、库存、配送、属性和变量标签。",
        product_edit,
        [
            "从 产品 > 全部产品 点击商品名称进入此页；先确认商品类型后，再填写常规价格或促销价。",
            "在 产品数据 > 库存 填写唯一 SKU；启用“管理库存”后填写数量和低库存阈值。",
            "产品图片与产品相册分别维护主图和补充图；更改后在前台商品页和购物车复核。",
        ],
        "示例页面中的字段仅用于说明位置。不得将生产密钥、客户隐私信息或测试数据当作正式商品资料发布。",
    )

    add_capture_page(
        doc,
        "7.1 商品库存：SKU、库存管理与售卖状态",
        "同一商品的库存页用于维护 SKU、是否按数量扣减，以及前台可售状态。以下截图来自当前后台的简单商品示例。",
        product_stock,
        [
            "在 1 号区域打开“库存”标签；SKU 是商品与订单、仓库或外部系统匹配时使用的唯一编号。",
            "2 号区域填写或核对 SKU。需要按实际库存扣减时，在 3 号区域启用“跟踪库存数量”。",
            "4 号区域选择有货、无货或延迟交货；修改后点击右侧“更新”，再到前台检查是否仍可购买。",
        ],
        "不要通过修改 SKU 来表示颜色、规格或批次。存在可独立销售的规格时，应使用属性和变体分别维护价格、SKU 与库存。",
    )

    add_capture_page(
        doc,
        "8. 报价记录：客户提交后的留存入口",
        "客户的 3D 询价会留存在 toKraft > 报价记录。列表可查看报价编号、客户、材料与数量、状态、负责人和模型文件。",
        quote_records,
        [
            "按状态和日期筛选报价；点击报价编号进入详情，不要只依赖邮件收件箱跟进。",
            "列表中的演示数据用于展示工作流；上线前应替换或删除演示询价，避免混入真实业务数据。",
            "新询价进入后先核对客户邮件、材料和数量，再分配负责人。",
        ],
    )

    add_capture_page(
        doc,
        "9. 报价详情：状态、负责人和正式报价",
        "报价详情页集中保存客户工艺参数、内部跟进状态、负责人、正式报价金额和内部备注。",
        quote_detail,
        [
            "核对材料、颜色、数量、填充率、壁厚、层高、支撑、附着方式、备注和模型文件。",
            "在“跟进状态”中更新待处理、已报价、已成交或未成交；同时指定负责人，保证交接可追溯。",
            "填写正式报价和内部跟进备注后点击更新；客户可见的报价应通过受控的业务沟通渠道发送。",
        ],
    )

    add_capture_page(
        doc,
        "10. 材料库：材料名称与产品关联",
        "toKraft > 材料库当前以分类法管理材料名称、别名、描述和关联产品数，用于保持产品与材料筛选的一致性。",
        materials,
        [
            "编辑既有材料可调整名称、别名与描述；删除前必须确认关联产品数为零或已完成迁移。",
            "在商品编辑页为商品选择正确的材料库项目，避免同一种材料建立多个近似名称。",
            "材料详情或前台文案需通过实际前台页面复核，不能仅根据后台分类名称判断展示结果。",
        ],
    )

    add_capture_page(
        doc,
        "11. 支付：付款提供程序的配置入口",
        "WooCommerce > 设置 > 付款是启用或管理收款方式的唯一入口。截图中的卡片加载状态也说明：上线前应确认每个网关实际加载完成。",
        payments,
        [
            "先确定业务实际支持的付款方式、适用地区和收款责任人，再启用对应支付提供程序。",
            "进入某一支付方式的管理页后填写客户可见标题、说明和必要的商户资料；密钥只保存于受控配置中。",
            "完成后必须使用测试订单验证付款成功、取消、邮件通知、订单状态和库存扣减。",
        ],
        "若支付提供程序长期停留在加载状态，请由技术人员检查 WooCommerce、浏览器控制台和插件兼容性；未确认前不要上线收款。",
    )

    add_capture_page(
        doc,
        "12. Blog 编辑页面：封面与发布前复核",
        "文章编辑器用于修改标题、正文、分类、标签与特色图片。发布前以统一封面尺寸和前台预览为准。",
        blog_edit,
        [
            "封面图使用统一的横向比例和清晰主题；每篇文章都检查标题、摘要、正文首屏和相关推荐的视觉效果。",
            "先保存草稿或预览，再更新已发布文章；不要直接覆盖尚未确认的客户内容。",
            "更新后打开前台文章页，检查封面裁切、标题叠加、正文图片和移动端布局。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "13. 页面、菜单与基础设置", 1)
    add_text(doc, "页面模块用于首页、Blog、Cart、Checkout 等固定页面；外观模块通常用于菜单和主题外观。此类修改会影响全站导航，建议先在测试环境验证。", size=10.5, color="000000")
    add_heading(doc, "页面维护", 2)
    for text in [
        "进入“页面”后，点击页面标题打开编辑器；更新首页、购物车或结账页前，先在前台确认该页面承担的功能。",
        "固定页面改动后，应分别检查桌面与移动端，并确认页脚、导航和按钮链接仍然可用。",
        "涉及页面模板或短代码的内容，不建议直接删除；先复制保留原文，再做最小范围调整。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "菜单与外观", 2)
    for text in [
        "菜单调整前记录原始顺序和链接；若有多语言、页脚菜单或移动端菜单，请分别确认。",
        "仅修改文字或链接时，优先使用菜单编辑；涉及布局、颜色或主题文件时，应交由管理员处理。",
    ]:
        add_bullet(doc, text)
    add_note_box(doc, "设置原则", "站点地址、固定链接、支付、税率、物流和插件配置属于系统级设置。非管理员不要修改；如确有需求，先做完整备份并记录变更项。", RED, "FDECEC")

    doc.add_page_break()
    add_heading(doc, "14. 发布前检查与常见问题", 1)
    add_heading(doc, "发布前 60 秒检查", 2)
    for text in [
        "标题、摘要、分类、标签和特色图是否完整；图片比例是否一致。",
        "产品的 SKU、价格、库存、分类和前台展示图是否正确。",
        "链接、按钮、购物车和结账页能否正常打开；不要只在已登录后台状态下判断。",
        "正式发布后，用无痕窗口或其他设备检查前台结果。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "常见问题处理", 2)
    problems = doc.add_table(rows=1, cols=2)
    set_table_geometry(problems, [2500, 6860])
    for cell, text in zip(problems.rows[0].cells, ["现象", "建议处理"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for left, right in [
        ("登录后跳回登录页", "不要直接改数据库地址。先确认当前访问域名、浏览器 Cookie 和 WordPress 地址设置，并联系管理员统一处理。"),
        ("产品图片不显示", "检查特色图片/产品图库是否已上传、图片链接是否有效，并在前台刷新缓存后复核。"),
        ("订单状态无法判断", "不要猜测付款状态；进入订单详情核对付款记录和订单备注，必要时由商城负责人确认。"),
        ("改动后前台未变化", "先清浏览器缓存并使用无痕窗口检查；仍无变化时记录页面、时间和操作步骤后提交给技术人员。"),
    ]:
        cells = problems.add_row().cells
        for cell, text in zip(cells, [left, right]):
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=10, color="000000")
    add_text(doc, "支持交接时请同时提供：操作账号角色、发生时间、相关页面 URL、截图、复现步骤以及是否影响订单或客户。", size=10.3, color=NAVY, bold=True, before=16, after=0)

    doc.add_page_break()
    add_heading(doc, "15. 首页运营：内容区块、素材与前台验证", 1)
    add_text(doc, "首页不是通过普通“页面编辑器”维护，而是由专门的 toKraft 首页内容区块管理器控制。该入口用于管理首页文案、按钮、图片和展示数量；材料、案例、设备等卡片内容则从各自的资料库自动读取。", size=10.5, color="000000")
    add_heading(doc, "进入路径与可编辑范围", 2)
    home_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(home_table, [2700, 3420, 3240])
    for cell, text in zip(home_table.rows[0].cells, ["后台位置", "可以维护", "维护后前台位置"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for row_data in [
        ("toKraft > 首页内容区块", "Hero 标题、说明、轮播图、两个主按钮及卖点", "首页首屏"),
        ("同一页面的服务/商店区块", "卡片标题、说明、要点、按钮和配图", "首页服务与商城入口"),
        ("同一页面的材料/案例区块", "区块标题、说明、显示数量、按钮链接", "首页材料与应用展示"),
        ("toKraft > 材料库 / 应用案例 / 设备", "卡片实际内容、图片、是否精选、排序", "首页自动读取的材料、案例和设备卡片"),
    ]:
        cells = home_table.add_row().cells
        for cell, text in zip(cells, row_data):
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=9.8, color="000000")
    add_heading(doc, "首页运营步骤", 2)
    for text in [
        "进入“toKraft > 首页内容区块”，先确认需要调整的是首屏、服务/商店、设备、材料/案例还是底部说明区块。不要把材料卡片的实际信息改在首页文案字段里。",
        "首屏若使用单图，可替换主视觉；若使用轮播模式，可选择最多 5 张轮播图。更新前统一准备裁切后的图片，并检查按钮文字和目标链接。",
        "服务与商店卡片中的“要点”按一行一项填写。按钮链接优先使用站内相对路径，例如 /quote/、/shop/、/materials/，避免误写临时域名。",
        "材料、案例和设备均会按后台数据自动显示。首页仅决定显示数量和区块文案；要让某个材料或案例上首页，需要到对应资料库勾选“Show on homepage / 首页展示”。",
        "保存后打开前台首页，用无痕窗口检查：首屏文字换行、按钮跳转、图片裁切、材料卡片数量和案例排序。",
    ]:
        add_bullet(doc, text)
    add_note_box(doc, "首页运营边界", "首页管理器不等同于拖拽式页面搭建器。涉及结构、组件样式或区块顺序的改动，应由技术人员处理；运营人员应限定在文案、图片、链接和展示数量。", GOLD, "FFF8E8")

    doc.add_page_break()
    add_heading(doc, "16. 商品维护与库存：从编辑到前台销售", 1)
    add_text(doc, "商品维护需要同时完成基础信息、价格、库存、图片和前台展示信息。仅改商品标题或库存状态，可能无法让客户得到正确的购买信息。", size=10.5, color="000000")
    add_heading(doc, "单个商品编辑路径", 2)
    for text in [
        "进入“产品 > 全部产品”，点击商品名称进入编辑页。先确认商品类型（简单产品或可变产品），再开始调整。",
        "在页面顶部填写商品名称、长描述和简短描述；右侧设置商品分类、产品主图与产品图库。主图建议 1600 x 1600 px，图库建议至少 1600 x 1200 px，保持光线与背景一致。",
        "打开“产品数据 > 常规”：维护常规售价、促销价及税务状态。价格改动后，必须在前台商品页和购物车各检查一次。",
        "打开“产品数据 > 库存”：填写唯一 SKU；若需要按件扣减，勾选“管理库存”，设置库存数量、低库存阈值和缺货状态。未启用库存管理时，至少要正确设置“有货 / 缺货 / 缺货待补”。",
        "打开“产品数据 > 配送”：填写真实重量、长宽高。页面右侧的“toKraft Product Setup”用于填写前台展示尺寸说明；它不替代配送尺寸。",
        "如商品有颜色、规格或材质变体，先在“属性”中建立属性，再进入“变体”创建每个可售组合，分别维护价格、SKU、库存和图片。",
        "点击“更新”后，用前台无痕窗口确认商品图片、价格、可选项、库存提示、加入购物车和结账流程。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "库存处理规则", 2)
    stock_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(stock_table, [2400, 3360, 3600])
    for cell, text in zip(stock_table.rows[0].cells, ["场景", "后台操作", "必须复核"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for row_data in [
        ("新增可售库存", "产品数据 > 库存 > 库存数量", "SKU 是否正确；前台是否显示“有货”"),
        ("暂停销售", "库存状态改为“缺货”，或将商品改为草稿", "前台不能继续加入购物车"),
        ("变体缺货", "编辑该变体自己的库存和状态", "其他变体不应被误设为缺货"),
        ("批量修改", "先筛选目标商品，再在批量编辑中修改", "筛选结果数量、价格单位和影响范围"),
    ]:
        cells = stock_table.add_row().cells
        for cell, text in zip(cells, row_data):
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=9.8, color="000000")
    add_note_box(doc, "库存提醒", "WooCommerce 只会按已启用的“管理库存”逻辑扣减数量。对变体商品，库存通常在变体层级维护；不要仅修改父商品的状态后就认为全部变体已更新。", RED, "FDECEC")

    doc.add_page_break()
    add_heading(doc, "17. 支付定义：启用前要准备什么", 1)
    add_text(doc, "支付方式在“WooCommerce > 设置 > 支付”中定义。支付开关不仅影响结账页，还会影响订单状态、客户邮件和履约开始时间，因此必须以实际收款能力为准。", size=10.5, color="000000")
    add_heading(doc, "当前环境的支付状态", 2)
    add_note_box(doc, "当前配置核对结果", "本次截图时，付款提供程序列表仍处于加载状态，无法据此确认任何网关已启用或已完成收款资料配置。上线前请逐项打开并核验实际可见的付款方式、商户资料、结账页显示和测试订单结果。", RED, "FDECEC")
    add_heading(doc, "每种支付方式需准备的资料", 2)
    pay_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(pay_table, [2340, 4020, 3000])
    for cell, text in zip(pay_table.rows[0].cells, ["方式", "启用前必须准备", "配置后验证"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for row_data in [
        ("银行转账（BACS）", "收款主体名称、银行名称、账号/IBAN、SWIFT/BIC、订单付款备注与到账确认流程", "结账页文案、客户邮件、订单是否进入待付款/保留状态"),
        ("货到付款（COD）", "可用配送区域、适用商品、收款责任人和拒收处理规则", "仅在允许的配送方式和地区显示"),
        ("支票付款", "收件人、邮寄地址、付款说明与人工核销规则", "客户看到的说明与内部核销流程一致"),
        ("PayPal 或在线卡支付", "已认证商户账号、生产环境 API/密钥、回调地址、退款和对账负责人", "沙盒测试成功后才切生产；完成付款、取消、退款各测一次"),
    ]:
        cells = pay_table.add_row().cells
        for cell, text in zip(cells, row_data):
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=9.6, color="000000")
    add_heading(doc, "配置和上线步骤", 2)
    for text in [
        "先由业务负责人确认实际可接受的付款方式和地区；不提供的方式请关闭，避免客户下单后无法收款。",
        "进入“WooCommerce > 设置 > 支付”，点击相应方式的管理按钮，填写标题、结账页说明、订单邮件说明和必要账户资料。",
        "银行转账、支票、货到付款是人工确认型方式：必须明确谁确认收款、何时将订单改为处理中或已完成。",
        "第三方在线支付必须先使用测试模式和测试订单验证；生产密钥不得写入页面文案、截图或普通文档。",
        "上线前在无痕窗口走完一次下单流程，检查付款方式显示、客户邮件、订单后台状态与库存扣减是否一致。",
    ]:
        add_bullet(doc, text)

    doc.add_page_break()
    add_heading(doc, "18. 3D 材料与报价表单：资料维护和提交数据去向", 1)
    add_text(doc, "3D 业务包含两条不同的数据链路：材料资料链路负责前台材料库、首页材料卡片和报价下拉选项；报价表单链路负责接收客户的模型参数与文件。二者需要分别维护和查看。", size=10.5, color="000000")
    add_heading(doc, "材料库：维护首页和报价页可选材料", 2)
    for text in [
        "进入“toKraft > 材料库”。当前后台页面维护材料名称、别名、描述和关联产品数；名称应与报价页和商品页使用的术语保持一致。",
        "编辑材料前先检查关联产品数；若需要合并或改名，先迁移关联商品，避免前台筛选页出现空分类或重复材料。",
        "商品的主图和图集在“产品 > 编辑产品”中维护，不要把材料分类页当作产品图片管理器。",
        "首页材料区块的标题、说明和展示数量在“toKraft > 首页内容区块”中维护；保存后必须在前台复核实际卡片。",
        "最终报价仍应在人工审查模型文件后确认，材料分类本身不等同于报价规则。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "3D 报价表单：客户提交后去哪里看", 2)
    add_note_box(doc, "当前系统行为", "客户提交 3D 询价后，系统会在“toKraft > 报价记录”中创建可筛选记录。详情页保存联系人、邮箱、公司、材料、颜色、数量、工艺参数、模型文件、跟进状态、负责人、正式报价和内部备注；通知邮件可作为辅助提醒，而不是唯一业务记录。", GOLD, "FFF8E8")
    for text in [
        "查看新报价：进入“toKraft > 报价记录”，按日期或状态筛选；点击报价编号进入详情页。",
        "查看模型文件：在报价详情的“模型文件”字段打开关联文件；必要时再通过“媒体 > 媒体库”按上传时间辅助核对。",
        "报价审核：核对材料、数量、层高、填充率、支撑、尺寸与可制造性；填写正式报价和内部备注，并明确分配负责人。",
        "团队协作：状态统一使用待处理、已报价、已成交或未成交；更新前后都保留关键决策和客户沟通摘要，避免只依赖个人邮箱。",
    ]:
        add_bullet(doc, text)
    add_note_box(doc, "交付前检查", "报价记录已具备状态、负责人、报价金额与备注字段。上线前请提交一条测试询价，验证记录生成、文件关联、邮件提醒和负责人交接是否完整；测试数据需明确标识或按规则清理。", RED, "FDECEC")

    doc.core_properties.title = "toKraft WordPress 后台使用指南"
    doc.core_properties.subject = "客户交付 - WordPress 与 WooCommerce 后台操作手册"
    doc.core_properties.author = "toKraft"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
